"""Document loaders for various file formats."""

from __future__ import annotations

import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import structlog

logger = structlog.get_logger(__name__)


@dataclass
class Document:
    """Represents a loaded document fragment with metadata."""

    content: str
    metadata: dict[str, Any] = field(default_factory=dict)

    def __post_init__(self) -> None:
        # Guarantee standard metadata keys exist
        self.metadata.setdefault("source", "")
        self.metadata.setdefault("page", None)
        self.metadata.setdefault("filename", "")
        self.metadata.setdefault("file_type", "")
        self.metadata.setdefault("section", None)


class BaseLoader(ABC):
    """Abstract base class for document loaders."""

    @abstractmethod
    def load(self, file_path: str | Path) -> list[Document]:
        """Load a file and return a list of Document objects.

        Args:
            file_path: Path to the file to load.

        Returns:
            List of Document objects extracted from the file.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file is empty or corrupt.
        """
        ...

    def _validate_path(self, file_path: str | Path) -> Path:
        """Validate that the file exists and is not empty."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        if not path.is_file():
            raise ValueError(f"Path is not a file: {path}")
        if path.stat().st_size == 0:
            raise ValueError(f"File is empty: {path}")
        return path

    def _base_metadata(self, path: Path) -> dict[str, Any]:
        """Build standard metadata dict from a file path."""
        return {
            "source": str(path.resolve()),
            "filename": path.name,
            "file_type": path.suffix.lstrip(".").lower(),
        }


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


class PDFLoader(BaseLoader):
    """Load PDF files using pypdf, one Document per page."""

    def load(self, file_path: str | Path) -> list[Document]:
        path = self._validate_path(file_path)
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ImportError(
                "pypdf is required for PDF loading. Install it: pip install pypdf"
            ) from exc

        try:
            reader = PdfReader(str(path))
        except Exception as exc:
            raise ValueError(f"Failed to read PDF {path}: {exc}") from exc

        if len(reader.pages) == 0:
            raise ValueError(f"PDF has no pages: {path}")

        documents: list[Document] = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                logger.debug("pdf.empty_page", file=path.name, page=page_num)
                continue
            meta = self._base_metadata(path)
            meta["page"] = page_num
            documents.append(Document(content=text, metadata=meta))

        if not documents:
            raise ValueError(f"PDF contains no extractable text: {path}")

        logger.info("pdf.loaded", file=path.name, pages=len(documents))
        return documents


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


class DOCXLoader(BaseLoader):
    """Load DOCX files using python-docx, one Document per non-empty paragraph."""

    def load(self, file_path: str | Path) -> list[Document]:
        path = self._validate_path(file_path)
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX loading. Install it: pip install python-docx"
            ) from exc

        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            raise ValueError(f"Failed to read DOCX {path}: {exc}") from exc

        documents: list[Document] = []
        current_section: str | None = None

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Track headings as section context
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                current_section = text

            meta = self._base_metadata(path)
            meta["page"] = None
            meta["section"] = current_section
            meta["paragraph_index"] = idx
            documents.append(Document(content=text, metadata=meta))

        if not documents:
            raise ValueError(f"DOCX contains no text: {path}")

        logger.info("docx.loaded", file=path.name, paragraphs=len(documents))
        return documents


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


class TextLoader(BaseLoader):
    """Load plain .txt files as a single Document."""

    def load(self, file_path: str | Path) -> list[Document]:
        path = self._validate_path(file_path)
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")

        text = text.strip()
        if not text:
            raise ValueError(f"Text file is empty: {path}")

        meta = self._base_metadata(path)
        logger.info("text.loaded", file=path.name, chars=len(text))
        return [Document(content=text, metadata=meta)]


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

_FRONTMATTER_RE = re.compile(r"^---\s*\n.*?\n---\s*\n", re.DOTALL)


class MarkdownLoader(BaseLoader):
    """Load Markdown files, stripping YAML frontmatter."""

    def load(self, file_path: str | Path) -> list[Document]:
        path = self._validate_path(file_path)
        try:
            raw = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            raw = path.read_text(encoding="latin-1")

        # Strip frontmatter
        text = _FRONTMATTER_RE.sub("", raw).strip()
        if not text:
            raise ValueError(f"Markdown file has no content after stripping frontmatter: {path}")

        meta = self._base_metadata(path)
        # Extract first heading as section
        heading_match = re.search(r"^#+\s+(.+)$", text, re.MULTILINE)
        if heading_match:
            meta["section"] = heading_match.group(1).strip()

        logger.info("markdown.loaded", file=path.name, chars=len(text))
        return [Document(content=text, metadata=meta)]


# ---------------------------------------------------------------------------
# Factory
# ---------------------------------------------------------------------------

_LOADER_MAP: dict[str, type[BaseLoader]] = {
    ".pdf": PDFLoader,
    ".docx": DOCXLoader,
    ".txt": TextLoader,
    ".md": MarkdownLoader,
    ".markdown": MarkdownLoader,
}


def get_loader(file_path: str | Path) -> BaseLoader:
    """Return an appropriate loader instance for *file_path* based on extension.

    Args:
        file_path: Path to the file.

    Returns:
        A BaseLoader subclass instance.

    Raises:
        ValueError: If the file extension is not supported.
    """
    ext = Path(file_path).suffix.lower()
    loader_cls = _LOADER_MAP.get(ext)
    if loader_cls is None:
        supported = ", ".join(sorted(_LOADER_MAP.keys()))
        raise ValueError(f"Unsupported file extension '{ext}'. Supported: {supported}")
    return loader_cls()
